from flask import Blueprint, render_template, request, redirect, url_for
import ast
import os
from collections import Counter
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from comtypes.client import CreateObject
import comtypes
from module.report_module import report_module

reportPDF_module = Blueprint("reportPDF_module", __name__)
@reportPDF_module.route("/reportPDF_result", methods=["POST"])

def reportPDF_result():
    class SecurityReport:
        def __init__(self, base_directory, save_path):
            self.base_directory = base_directory
            self.save_path = save_path
            self.document = Document()

        def split_name(self, directory):
            dict_names = []
            try:
                for data_path in os.listdir(directory):
                    with open(f'{directory}/{data_path}', 'r', encoding='utf-8') as file:
                        data = file.read()
                        rdata = str(data).split('$')
                        who = rdata[0]
                        rdata = rdata[1]
                        data_list = ast.literal_eval(rdata)
            except:
                pass
            return who, data_list

        def split_data(self, directory):
            keys = []
            values = []
            try:
                for data_path in os.listdir(directory):
                    with open(directory+'\\'+data_path, 'r', encoding='utf-8') as file:
                        data = file.read()
                        data_list = ast.literal_eval(data)
                        for data_count in range(len(data_list.keys())):
                            keys.append(list(data_list.keys())[data_count])
                            values.append(list(data_list.values())[data_count])
            except:
                pass
            return keys, values


        def create_table(self, who, content):
            self.document.add_heading('기업 보안 보고서', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.document.add_heading('github', level=2)

            rcontent = []
            rpath = []
            repo = []
            for i in range(len(content.values())):
                for j in range(len(list(content.values())[i])):
                    repo.append(list(content.keys())[i])
                    rpath.append((list(content.values())[i][j]['path']))
                    rcontent.append((list(content.values())[i][j])['content'])

            table = self.document.add_table(rows=len(rcontent) + 1, cols=4)
            table.style = self.document.styles['Table Grid']

            table.cell(0, 0).text = 'who'
            table.cell(0, 1).text = 'repository'
            table.cell(0, 2).text = 'path'
            table.cell(0, 3).text = 'content'

            start_row_repo = 1  # Variable initialization
            start_row_path = 1  # Variable initialization

            for row in range(1, len(rcontent) + 1):  
                for col in range(4):
                    if row == 0:
                        continue
                    table.cell(row, col).text = [who, repo[row - 1], rpath[row - 1], rcontent[row - 1]][col]

                # Who cell merge
                if row == len(rcontent):
                    merged_cell = table.cell(1, 0).merge(table.cell(len(rcontent), 0))
                    merged_cell.text = who
                    merged_cell = self.merge_style(merged_cell)

                # Repository cell merge
                if row < len(rcontent) and repo[row] != repo[row - 1]:
                    merged_cell = table.cell(start_row_repo, 1).merge(table.cell(row, 1))
                    merged_cell.text = repo[start_row_repo - 1]
                    merged_cell = self.merge_style(merged_cell)
                    start_row_repo = row + 1

                # Path cell merge
                if row < len(rcontent) and (repo[row], rpath[row]) != (repo[row - 1], rpath[row - 1]):
                    merged_cell = table.cell(start_row_path, 2).merge(table.cell(row, 2))
                    merged_cell.text = rpath[start_row_path - 1]
                    merged_cell = self.merge_style(merged_cell)
                    start_row_path = row + 1

        def create_domain_table(self, keys, values):
            self.document.add_heading('domain', level=2)

            keywords = []
            emails = []
            phones = []
            filter_keyword = []

            for i in range(len(values)):
                keywords.append(values[i][0]['keywords'])
                emails.append(values[i][0]['emails'])
                phones.append(values[i][0]['phones'])
                filter_keyword.append(values[i][0]['filter_keyword'])

            table = self.document.add_table(rows=len(keys)*3 +1 , cols=4)
            table.style = self.document.styles['Table Grid']

            table.cell(0, 0).text = 'URL'
            table.cell(0, 1).text = 'filter_keyword'
            table.cell(0, 2).text = 'emails'
            table.cell(0, 3).text = 'phones'

            for numcount in range(len(keys)):
                for row in range(len(keys)*3 +1):
                    if row == 0:
                        pass
                    else:
                        for col in range(4):
                            if (row - 2) % 3 == 0 or row % 3 == 0:
                                merged_cell = table.cell(row, 0).merge(table.cell(row,3))
                                merged_cell.text = ["keyword", str(keywords[numcount])][row % 3 == 0]
                                merged_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
                                merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            else:
                                table.cell(row, col).text = [keys[numcount], filter_keyword[numcount], str(emails[numcount]), str(phones[numcount])][col]

        def create_network_table(self, directory):
            self.document.add_heading('network', level=2)

            data_list = []
            try:
                for data_path in os.listdir(directory):
                    with open(directory + '\\' + data_path, 'r', encoding='utf-8') as file:
                        data = file.read()
                        data_list.append(ast.literal_eval(data))
            except:
                pass

            table = self.document.add_table(rows=len(data_list) + 1, cols=5)
            table.style = self.document.styles['Table Grid']

            headers = ['ip', 'hostname', 'open port', 'server version', 'CVE']
            for idx, h in enumerate(headers):
                table.cell(0, idx).text = h

            for row in range(1, len(data_list) + 1):
                data = data_list[row-1]
                for col in range(5):
                    if col == 0:
                        table.cell(row, col).text = data['nmap_result']['ip_address']
                    elif col == 1:
                        table.cell(row, col).text = data['nmap_result']['rdns_records']
                    elif col == 2:
                        table.cell(row, col).text = str(data['nmap_result']['port_info'])
                    elif col == 3:
                        table.cell(row, col).text = data.get('server_info', 'NONE')  # If server_info is not present
                    elif col == 4:
                        cve_collection = []
                        for item_list in data['cve_info']:
                            for item in item_list:
                                cve_collection.append(item.get('CVE ID', 'NONE'))  # If CVE ID is not present
                        table.cell(row, col).text = str(cve_collection)

        def save_as_pdf(self):
            # Create a Word application object
            word = CreateObject("Word.Application")
            # Show the Word application
            word.Visible = False
            # Load the document
            in_file = os.path.abspath(self.save_path)
            out_file = os.path.abspath(self.save_path.replace(".docx", ".pdf"))
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=17)  # FileFormat=17 is for .pdf
            doc.Close()
            word.Quit()

            # remove the original .docx file
            if os.path.exists(in_file):
                os.remove(in_file)

        def save(self):
            self.document.save(self.save_path)

        @staticmethod
        def merge_style(merged_cell):
            merged_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            return merged_cell

        def generate_report(self):
            modules = ['github', 'domain', 'network']
            methods = [self.create_table, self.create_domain_table, self.create_network_table]

            for module, method in zip(modules, methods):
                directory = os.path.join(self.base_directory, f'{module}_module')
                if module == 'github':
                    who, content = self.split_name(directory)
                    method(who, content)
                elif module == 'network':
                    method(directory)  # create_network_table only requires directory as input
                else:
                    keys, values = self.split_data(directory)
                    method(keys, values)

            self.save()
            self.save_as_pdf()

    comtypes.CoInitialize()

    if request.cookies.get('folder') is not None and request.cookies.get('folder') != '' :
        log_path = './crawling_log/' + request.cookies.get('folder').encode('latin-1').decode('utf-8') + '/'
    else:
        log_path = './crawling_log/none/'

    # Usage
    report = SecurityReport(
        base_directory=log_path,
        save_path=log_path + 'report_pdf.docx'
    )
    report.generate_report()

    report.save()
    report.save_as_pdf()   
    
    return render_template('reportPDF_result.html', Path=report.save_path)